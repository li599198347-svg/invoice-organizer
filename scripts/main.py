#!/usr/bin/env python3
"""差旅发票整理 - 按邮件分组处理发票和行程单"""
import os, sys, re, zipfile, shutil, imaplib, email, fitz, quopri
from email.header import decode_header
from email.utils import parsedate_to_datetime
from docx import Document
from docx.shared import Mm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from pdfminer.high_level import extract_text

SKILL_DIR = os.path.dirname(os.path.abspath(__file__))
CONFIG_FILE = os.path.join(os.path.dirname(SKILL_DIR), 'config.env')
OUTPUT_DIR = os.path.join(os.path.expanduser('~/Desktop'), '差旅发票')
TRIP_SHEET_DIR = os.path.join(OUTPUT_DIR, '行程单 PDF')

def load_config():
    memory_path = os.path.join(os.path.expanduser('~'), '.openclaw', 'workspace', 'MEMORY.md')
    cfg = {}
    if os.path.exists(memory_path):
        with open(memory_path, encoding='utf-8') as f:
            for match in re.finditer(r'^([A-Z_]+)=(.*)$', f.read(), re.MULTILINE):
                k, v = match.group(1), match.group(2).strip()
                if k in ('IMAP_HOST','IMAP_PORT','IMAP_USER','IMAP_PASS','IMAP_MAILBOX'):
                    cfg[k] = v
        print(f"  从记忆文件读取 IMAP 配置：{cfg.get('IMAP_USER','')}")
    if not cfg.get('IMAP_USER'):
        if os.path.exists(CONFIG_FILE):
            with open(CONFIG_FILE, encoding='utf-8') as f:
                for line in f:
                    if '=' in line and not line.startswith('#'):
                        k, v = line.split('=', 1)
                        if k.strip() in ('IMAP_HOST','IMAP_PORT','IMAP_USER','IMAP_PASS','IMAP_MAILBOX'):
                            cfg[k.strip()] = v.strip()
    if not cfg.get('IMAP_USER') or not cfg.get('IMAP_PASS'):
        print("❌ 未配置 IMAP 信息。"); sys.exit(1)
    return cfg

def decode_str(s):
    if s is None: return ''
    parts = decode_header(s)
    return ''.join([b.decode(cs or 'utf-8', errors='replace') if isinstance(b, bytes) else b for b, cs in parts])

def scan_invoices(cfg, start_date, end_date):
    print(f"\n📧 连接邮箱 {cfg['IMAP_USER']} ...")
    mail = imaplib.IMAP4_SSL(cfg.get('IMAP_HOST','imap.qq.com'), int(cfg.get('IMAP_PORT') or '993'))
    mail.login(cfg['IMAP_USER'], cfg['IMAP_PASS'].encode('ascii', 'replace').decode('ascii'))
    mail.select(cfg.get('IMAP_MAILBOX', 'INBOX'))
    _, msgs = mail.search(None, 'SINCE', f'"{start_date}"', 'BEFORE', f'"{end_date}"')
    seq_nums = [b.decode() for b in msgs[0].split()]
    print(f"  找到 {len(seq_nums)} 封邮件")
    invoice_emails = {}
    for seq in seq_nums:
        status, data = mail.fetch(seq.encode(), '(BODY[HEADER.FIELDS (DATE FROM SUBJECT)])')
        if status != 'OK' or not data or not data[0]: continue
        msg = email.message_from_bytes(data[0][1])
        sender = decode_str(msg.get('From',''))
        subject = decode_str(msg.get('Subject',''))
        try: email_date = parsedate_to_datetime(msg.get('Date','')).strftime('%Y-%m-%d')
        except: email_date = ''
        inv_type = None
        if 'xiaojukeji.com' in sender or 'didichuxing' in sender.lower():
            inv_type = '阳光出行' if '第三方' in subject else '滴滴'
        elif 'txffp.com' in sender or '通行费' in subject:
            inv_type = '通行费'
        elif '携程' in subject or 'ctrip' in sender.lower() or 'trip.com' in sender.lower():
            inv_type = '携程机票'  # ⚠️ 未经测试
        elif '火车票' in subject or '高铁' in subject or '铁路' in subject or '12306' in sender.lower():
            inv_type = '12306 火车票'
        elif '机票' in subject or '航空' in subject or '行程单' in subject or '客票' in subject:
            inv_type = '飞机票'  # ⚠️ 未经测试
        if inv_type:
            invoice_emails.setdefault(inv_type, []).append({'seq': seq, 'subject': subject, 'date': email_date})
            marker = '⚠️' if inv_type in ('火车票', '飞机票', '携程机票') else ''
            print(f"  [{seq}] {inv_type}: {subject[:50]} {marker}")
    mail.logout()
    return invoice_emails

def extract_trips_from_pdf(trip_path):
    """从行程单 PDF 提取所有行程明细"""
    try:
        raw_text = extract_text(trip_path)
        lines = raw_text.split('\n')
        non_empty = [line.strip() for line in lines if line.strip()]
        trips = []
        # 从行程单提取实际行程日期
        trip_date = ''
        # 滴滴/阳光出行格式：行程起止日期：2026-03-15 至 2026-03-15
        date_m = re.search(r'行程起止日期.*?(\d{4}-\d{2}-\d{2})', raw_text)
        if date_m: trip_date = date_m.group(1)
        # 通行费格式：通行日期 起止 \n 1 \n 20260310 \n 20260310（用 DOTALL 匹配换行）
        if not trip_date:
            date_m = re.search(r'通行日期.*?(\d{8})', raw_text, re.DOTALL)
            if date_m:
                d = date_m.group(1)
                trip_date = f"{d[:4]}-{d[4:6]}-{d[6:8]}"
        
        # 找到所有数字里程行（第一个是里程，跳过金额）
        mileage_found = False
        for i, text in enumerate(non_empty):
            if re.match(r'^\d+\.\d+$', text):
                if mileage_found: continue  # 跳过金额行
                mileage_found = True
                
                seq_idx = next((k for k in range(i - 1, max(0, i - 20), -1) if non_empty[k].isdigit()), -1)
                if seq_idx < 0: continue
                between = non_empty[seq_idx + 1:i]
                
                # 特殊处理：如果序号后第一行包含多个字段（空格分隔的单行格式）
                if between and ' ' in between[0]:
                    # 单行格式："滴滴特快 03-13 16:22 周五 深圳市"
                    parts = between[0].split()
                    # 查找城市
                    cities = ['北京', '上海', '广州', '深圳', '成都', '重庆', '杭州', '西安', '南京', '武汉', '苏州', '天津']
                    city_idx = next((k for k, t in enumerate(parts) if t in cities or (len(t) <= 5 and '市' in t)), -1)
                    # 起点终点可能在后续行
                    from_to_parts = parts[city_idx + 1:] if city_idx >= 0 else []
                    from_to_parts.extend(between[1:])  # 添加后续行
                    if len(from_to_parts) >= 2:
                        mid = len(from_to_parts) // 2
                        from_loc = ''.join(from_to_parts[:mid]).strip()
                        to_loc = ''.join(from_to_parts[mid:]).strip()
                        if from_loc and to_loc and len(from_loc) > 2 and len(to_loc) > 2:
                            trips.append({'from': from_loc, 'to': to_loc, 'date': trip_date})
                    continue
                
                if len(between) < 4: continue
                
                # 查找城市（常见城市名或含"市"的短词）
                cities = ['北京', '上海', '广州', '深圳', '成都', '重庆', '杭州', '西安', '南京', '武汉', '苏州', '天津']
                city_idx = next((k for k, t in enumerate(between) if t in cities or (len(t) <= 5 and '市' in t and len(t) >= 2)), -1)
                from_to = between[city_idx + 1:] if city_idx >= 0 else between
                if len(from_to) >= 2:
                    mid = len(from_to) // 2
                    from_loc = ''.join(from_to[:mid]).strip()
                    to_loc = ''.join(from_to[mid:]).strip()
                    from_loc = from_loc.replace('市', '').strip() if '市' in from_loc and len(from_loc) > 3 else from_loc
                    if from_loc and to_loc and len(from_loc) > 2 and len(to_loc) > 2 and not from_loc[0].isdigit() and '里程' not in from_loc and not re.search(r'\d+\.\d+', to_loc):
                        trips.append({'from': from_loc, 'to': to_loc, 'date': trip_date})
        return trips
    except Exception as e:
        print(f"    行程单提取失败：{e}")
        return []

def download_and_process(cfg, invoice_emails):
    print(f"\n📥 下载并处理附件...")
    mail = imaplib.IMAP4_SSL(cfg.get('IMAP_HOST','imap.qq.com'), int(cfg.get('IMAP_PORT') or '993'))
    mail.login(cfg['IMAP_USER'], cfg['IMAP_PASS'].encode('ascii', 'replace').decode('ascii'))
    mail.select(cfg.get('IMAP_MAILBOX', 'INBOX'))
    
    tmp_dir, pdf_dir = f'{OUTPUT_DIR}/_tmp', f'{OUTPUT_DIR}/发票 PDF'
    for d in [tmp_dir, pdf_dir, TRIP_SHEET_DIR]: os.makedirs(d, exist_ok=True)
    
    all_records = []
    
    for inv_type, emails in invoice_emails.items():
        print(f"\n  === {inv_type} ({len(emails)}封邮件) ===")
        for em in emails:
            status, data = mail.fetch(em['seq'].encode(), '(RFC822)')
            if status != 'OK' or not data or not data[0]: continue
            raw = data[0][1]
            msg = email.message_from_bytes(raw)
            
            # 携程机票：从邮件正文提取行程信息（默认当年，日期 + 行程放一起）
            ctrip_route = ''
            if inv_type == '携程机票':
                body = ''
                for part in msg.walk():
                    ctype = part.get_content_type()
                    if ctype == 'text/html':
                        try:
                            payload = part.get_payload()
                            if payload:
                                body += quopri.decodestring(payload).decode('utf-8', errors='replace')
                        except: pass
                
                # 简化正则：匹配"年 X 月 X 日 城市 - 城市"
                trip_m = re.search(r'年 (\d{1,2}) 月 (\d{1,2}) 日\s*([^\s<]+?)\s*[-–—]\s*([^\s<]+)', body)
                if trip_m:
                    # 日期 + 行程放一起
                    ctrip_route = f"2026-{trip_m.group(1).zfill(2)}-{trip_m.group(2).zfill(2)} {trip_m.group(3).strip()}-{trip_m.group(4).strip()}"
            
            counter, inv_infos, trips = {}, [], []
            for part in msg.walk():
                filename = decode_str(part.get_filename(''))
                if not filename: continue
                raw_data = part.get_payload(decode=True)
                if raw_data is None: continue
                
                if filename.endswith('.pdf'):
                    counter[filename] = counter.get(filename, 0) + 1
                    tmp_path = f"{tmp_dir}/{em['seq']}_{counter[filename]}_{filename}"
                    with open(tmp_path, 'wb') as f: f.write(raw_data)
                    
                    # 判断是行程单还是发票
                    is_trip_sheet = '行程' in filename or '报销单' in filename
                    is_toll_summary = '通行费' in filename and '汇总' in filename and '行程' not in filename
                    
                    if is_trip_sheet:
                        # 行程单
                        trip_saved = f"{TRIP_SHEET_DIR}/{em['seq']}_{filename}"
                        shutil.copy(tmp_path, trip_saved)
                        print(f"    [行程单] {filename}")
                        extracted = extract_trips_from_pdf(tmp_path)
                        for t in extracted: t['trip_sheet_path'] = trip_saved
                        trips.extend(extracted)
                        print(f"      提取 {len(extracted)} 个行程")
                    elif not is_toll_summary:
                        # 普通发票（跳过通行费汇总单）
                        try:
                            raw_text = extract_text(tmp_path)
                            # 彻底去掉所有空白字符（包括特殊空白如\u00a0）
                            text = raw_text.replace(' ', '').replace('\u00a0', '').replace('\t', '').replace('\n', '').replace('\r', '')
                        except: text = ''
                        inv_no = ''
                        from_loc, to_loc = '', ''
                        
                        # 提取发票号和行程信息
                        inv_no = ''
                        from_loc, to_loc = '', ''
                        trip_date = ''
                        
                        if inv_type in ('滴滴', '阳光出行'):
                            # 提取发票号
                            m = re.search(r'发票号码.*?([\d]+)', text)
                            if m: inv_no = m.group(1)
                            # 提取行程日期
                            date_m = re.search(r'出行日期 [：:]+\s*(\d{4}[-/.]\d{1,2}[-/.]\d{1,2})', text)
                            if date_m: trip_date = date_m.group(1).replace('/', '-')
                            # 提取起点终点
                            from_m = re.search(r'出\s*发\s*地 [：:]*\s*([\u4e00-\u9fffA-Za-z0-9\(\)\|\-]{2,50}?)', text)
                            to_m = re.search(r'到\s*达\s*地 [：:]*\s*([\u4e00-\u9fffA-Za-z0-9\(\)\|\-]{2,50}?)', text)
                            if from_m: from_loc = from_m.group(1).strip()
                            if to_m: to_loc = to_m.group(1).strip()
                            
                        elif inv_type == '12306 火车票':
                            # 12306 格式特殊：文字连续无换行
                            # 发票号：发票号码:26329166827000948449
                            m = re.search(r'发票号码 [：:]?\s*(\d{20})', text)
                            if m: inv_no = m.group(1)
                            
                            # 车次 + 日期：G83632026 年 02 月 23 日（车次和年份可能连在一起）
                            # 车次号通常 1-4 位数字，年份 4 位数字
                            # \s*允许数字和"年/月/日"之间有无空格
                            train_m = re.search(r'([GDKZTC]\d{1,4})(\d{4})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})\s*日', text)
                            if train_m:
                                trip_date = f"{train_m.group(2)}-{train_m.group(3).zfill(2)}-{train_m.group(4).zfill(2)}"
                            
                            # 站点：徐州东站、上海虹桥站
                            # 12306 PDF 中站点顺序：到达站在前，出发站在后（因为文字是从右到左读的）
                            stations = []
                            for m in re.finditer(r'([北京上海广州深圳徐州南京杭州武汉成都重庆西安苏州天津虹桥][^\s:：]*?站)', text):
                                station = m.group(1)
                                if '站' in station and '开票' not in station and '电子客票' not in station:
                                    stations.append(station)
                            # 反转顺序：第一个是出发站，第二个是到达站
                            if len(stations) >= 2:
                                from_loc = stations[-1].strip()  # 最后一个（徐州东站）
                                to_loc = stations[-2].strip()   # 倒数第二个（上海虹桥站）
                            elif len(stations) == 1:
                                from_loc = stations[0].strip()
                            
                            # 票价：票价:￥306.00
                            price_m = re.search(r'票价 [：:]?\s*[￥¥]?(\d+\.?\d*)', text)
                            if price_m:
                                amount = float(price_m.group(1))
                            
                        elif inv_type == '飞机票':  # ⚠️ 未经测试
                            m = re.search(r'发票代码 [：:]\s*([\d]+)', text) or re.search(r'发票号码 [：:]\s*([\d]+)', text)
                            if m: inv_no = m.group(1)
                            # 提取出发地/目的地
                            from_m = re.search(r'出发地 [：:]\s*(.+?)$', text, re.M) or re.search(r'始发站 [：:]\s*(.+?)$', text, re.M)
                            to_m = re.search(r'目的地 [：:]\s*(.+?)$', text, re.M) or re.search(r'到达站 [：:]\s*(.+?)$', text, re.M)
                            if from_m: from_loc = from_m.group(1).strip()
                            if to_m: to_loc = to_m.group(1).strip()
                            
                        elif inv_type == '携程机票':
                            # 简化处理：只提取发票号和金额
                            order_m = re.search(r'订单.*?(\d+)', os.path.basename(tmp_path))
                            if order_m: inv_no = f"携程{order_m.group(1)}"
                            # 携程不提取行程，显示"详见发票"
                            from_loc = '详见发票'
                            to_loc = ''
                        
                        # 如果没提取到发票号，用文件名
                        if not inv_no:
                            parts = os.path.basename(tmp_path).replace('.pdf','').split('_')
                            inv_no = parts[-1] if len(parts) > 2 else os.path.basename(tmp_path).replace('.pdf','')
                        
                        # PDF 以发票号命名
                        final_path = f"{pdf_dir}/{inv_no}.pdf"
                        if not os.path.exists(final_path): shutil.copy(tmp_path, final_path)
                        
                        # 提取金额
                        amounts = re.findall(r'(\d+\.\d+)', text)
                        amount = float(amounts[-1]) if amounts else 0
                        
                        inv_infos.append({
                            'inv_no': inv_no, 'amount': amount, 'pdf_path': final_path,
                            'type': inv_type, 'from': from_loc, 'to': to_loc, 'trip_date': trip_date
                        })
                        marker = '⚠️' if inv_type in ('火车票', '飞机票', '携程机票') else ''
                        print(f"    [发票] {inv_no} ¥{amount} {from_loc or ''}→{to_loc or ''} {marker}".strip())
                
                elif filename.endswith('.zip') and '通行费' in filename:
                    zip_path = f"{tmp_dir}/{em['seq']}_{filename}"
                    with open(zip_path, 'wb') as f: f.write(raw_data)
                    try:
                        with zipfile.ZipFile(zip_path, 'r') as zf:
                            xml_files = [n for n in zf.namelist() if n.startswith('xml/') and not n.endswith('/')]
                            pdf_files = {n.replace('pdf/','').replace('.pdf',''): n for n in zf.namelist() if n.startswith('pdf/')}
                            total_amt = 0
                            inv_nos = []
                            from_loc, to_loc = '', ''
                            for xml_name in sorted(xml_files):
                                if not xml_name.endswith('.xml'): continue
                                content = zf.read(xml_name).decode('utf-8', errors='replace')
                                inv_m = re.search(r'<EIid>(\d{20,})</EIid>', content) or re.search(r'<InvoiceNumber>(\d{20,})</InvoiceNumber>', content)
                                if inv_m:
                                    inv_no = inv_m.group(1).strip()
                                    inv_nos.append(inv_no)
                                    amt_m = re.search(r'<TotaltaxIncludedAmount>(\d+\.?\d*)</TotaltaxIncludedAmount>', content)
                                    if amt_m: total_amt += float(amt_m.group(1))
                                    # 从 XML 提取路线信息
                                    route_m = re.search(r'<TollRoad>([^<]+)</TollRoad>', content)
                                    if route_m and not from_loc:
                                        route_text = route_m.group(1)
                                        # 尝试分割起点终点
                                        if '→' in route_text:
                                            parts = route_text.split('→')
                                            from_loc, to_loc = parts[0].strip(), parts[-1].strip()
                                        elif '-' in route_text:
                                            parts = route_text.split('-')
                                            from_loc, to_loc = parts[0].strip(), parts[-1].strip()
                                    xml_key = xml_name.replace('xml/', '').replace('.xml', '')
                                    pdf_name = pdf_files.get(xml_key, '')
                                    if pdf_name:
                                        final_path = f"{pdf_dir}/{inv_no}.pdf"
                                        with open(final_path, 'wb') as f: f.write(zf.read(pdf_name))
                            # 保存所有子发票的 PDF 路径
                            pdf_paths = [f"{pdf_dir}/{inv_no}.pdf" for inv_no in inv_nos if inv_no]
                            inv_infos.append({
                                'inv_no': '\n'.join(inv_nos[:3]),
                                'amount': total_amt,
                                'pdf_path': pdf_paths[0] if pdf_paths else '',
                                'pdf_paths': pdf_paths,  # 所有子发票 PDF
                                'type': '通行费',
                                'from': from_loc or '详见发票',
                                'to': to_loc or '详见发票'
                            })
                            print(f"    [通行费] {len(inv_nos)}张 ¥{total_amt} {from_loc or '?'}→{to_loc or '?'}")
                    except Exception as e: print(f"    ZIP 处理错误：{e}")
            
            # 合并：多张发票对应多张行程单
            if inv_infos:
                # 滴滴/阳光出行：按顺序匹配行程单（一张发票对应一张行程单）
                if inv_type in ('滴滴', '阳光出行') and trips:
                    for i, inv in enumerate(inv_infos):
                        inv_pdf_paths = inv.get('pdf_paths', [inv['pdf_path']])
                        # 如果发票没有行程信息，从行程单中获取
                        if (not inv.get('from') or not inv.get('to')) and i < len(trips):
                            trip = trips[i]
                            trip_date = inv.get('trip_date') or trip.get('date', '')
                            all_records.append({
                                **inv, 'from': trip['from'], 'to': trip['to'],
                                'trip_date': trip_date, 'trip_count': 0, 'pdf_paths': inv_pdf_paths
                            })
                        else:
                            all_records.append({**inv, 'trip_count': 0, 'pdf_paths': inv_pdf_paths})
                # 通行费：从行程单获取入口站/出口站和日期
                elif inv_type == '通行费' and trips:
                    trip = trips[0]
                    trip_date = trip.get('date', '')
                    print(f"    通行费 trips[0].date={trip_date}, from={trip.get('from')}, to={trip.get('to')}")
                    for inv in inv_infos:
                        inv_pdf_paths = inv.get('pdf_paths', [inv['pdf_path']])
                        all_records.append({
                            **inv, 'from': trip['from'], 'to': trip['to'],
                            'trip_date': trip_date, 'trip_count': 0, 'pdf_paths': inv_pdf_paths
                        })
                # 通行费无行程单
                elif inv_type == '通行费':
                    for inv in inv_infos:
                        inv_pdf_paths = inv.get('pdf_paths', [inv['pdf_path']])
                        all_records.append({**inv, 'trip_count': 0, 'pdf_paths': inv_pdf_paths})
                # 携程简化处理
                elif inv_type == '携程机票':
                    for inv in inv_infos:
                        inv_pdf_paths = inv.get('pdf_paths', [inv['pdf_path']])
                        all_records.append({**inv, 'trip_count': 0, 'pdf_paths': inv_pdf_paths})
                # 12306 火车票保留提取的行程信息
                elif inv_type == '12306 火车票':
                    for inv in inv_infos:
                        inv_pdf_paths = inv.get('pdf_paths', [inv['pdf_path']])
                        all_records.append({**inv, 'trip_count': 0, 'pdf_paths': inv_pdf_paths})
                # 其他情况
                else:
                    for inv in inv_infos:
                        inv_pdf_paths = inv.get('pdf_paths', [inv['pdf_path']])
                        if trips:
                            if len(trips) == 1:
                                trip_date = inv.get('trip_date') or trips[0].get('date', '')
                                all_records.append({**inv, 'from': trips[0]['from'], 'to': trips[0]['to'], 'trip_date': trip_date, 'trip_count': 1, 'pdf_paths': inv_pdf_paths})
                            else:
                                trip_date = inv.get('trip_date') or trips[0].get('date', '')
                                all_records.append({**inv, 'from': f"对应{len(trips)}个行程", 'to': '', 'trip_date': trip_date, 'trip_count': len(trips), 'pdf_paths': inv_pdf_paths})
                        else:
                            all_records.append({**inv, 'from': inv.get('from', '详见发票'), 'to': inv.get('to', '详见发票'), 'trip_date': inv.get('trip_date', ''), 'trip_count': 0, 'pdf_paths': inv_pdf_paths})
            else:
                # 只有行程单没有发票
                for t in trips:
                    all_records.append({'inv_no': '待匹配', 'amount': 0, 'type': '行程单', 'from': t['from'], 'to': t['to'], 'trip_date': t.get('date', ''), 'trip_count': 1, 'pdf_path': '', 'pdf_paths': []})
    
    mail.logout()
    shutil.rmtree(tmp_dir, ignore_errors=True)
    return all_records

def build_word(all_records):
    print(f"\n📄 生成 Word...")
    
    # 预计算所有 PDF 页面
    all_pages = []
    for rec_idx, rec in enumerate(all_records):
        # 获取所有 PDF 路径（通行费有多张，其他只有一张）
        pdf_paths = rec.get('pdf_paths', [])
        if not pdf_paths and rec.get('pdf_path'):
            pdf_paths = [rec['pdf_path']]
        
        # 防错检查：如果没有 PDF 路径，记录警告
        if not pdf_paths:
            print(f"  ⚠️  记录{rec_idx+1} ({rec.get('type', '')} {rec.get('inv_no', '')[:20]}) 没有 PDF 路径")
            continue
        
        for pdf_idx, pdf_path in enumerate(pdf_paths):
            if not os.path.exists(pdf_path):
                print(f"  ⚠️  PDF 不存在：{pdf_path}")
                continue
            
            inv_no = rec.get('inv_no', '').split('\n')[0] if rec.get('inv_no') else ''
            suffix = f" ({pdf_idx+1}/{len(pdf_paths)})" if len(pdf_paths) > 1 else ''
            all_pages.append({
                'pdf_path': pdf_path,
                'page_text': f"{inv_no[:20]}{suffix}"
            })
    
    print(f"  共{len(all_pages)}张发票需要渲染")
    
    # 生成 Word
    doc = Document()
    for i, page_info in enumerate(all_pages):
        if i > 0:
            doc.add_section()
        
        sec = doc.sections[-1]
        sec.page_width, sec.page_height = Mm(210), Mm(148)
        sec.left_margin = sec.right_margin = sec.top_margin = sec.bottom_margin = Mm(8)
        
        try:
            pdf_doc = fitz.open(page_info['pdf_path'])
            pix = pdf_doc[0].get_pixmap(matrix=fitz.Matrix(2.0, 2.0))
            img_path = f'/tmp/travel_inv_{i+1:03d}.png'
            pix.save(img_path)
            pdf_doc.close()
            
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(img_path, width=Mm(194))
            
            p2 = doc.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run2 = p2.add_run()
            run2.text, run2.font.size, run2.font.name = f"Page {i+1} - {page_info['page_text']}", Pt(7), "Arial"
            
            # 更新记录的页码
            for rec in all_records:
                if page_info['pdf_path'] in rec.get('pdf_paths', []) or page_info['pdf_path'] == rec.get('pdf_path'):
                    if 'word_page' not in rec:
                        rec['word_page'] = i + 1
        except Exception as e:
            print(f"    PDF 渲染失败 {page_info['pdf_path']}: {e}")
    
    # 更新记录的页码（用于 Excel 关联）
    page_num = 1
    for rec in all_records:
        pdf_paths = rec.get('pdf_paths', [])
        if not pdf_paths and rec.get('pdf_path'):
            pdf_paths = [rec['pdf_path']]
        rec['word_page'] = page_num
        page_num += len(pdf_paths) if pdf_paths else 1
    
    out = f'{OUTPUT_DIR}/差旅发票汇总打印.docx'
    doc.save(out)
    print(f"  已保存：{out} (共{page_num-1}页)")
    return out

def build_excel(all_records):
    print(f"\n📊 生成 Excel...")
    wb = Workbook()
    ws = wb.active
    ws.title = "差旅发票汇总"
    thin = Border(left=Side(style='thin'),right=Side(style='thin'),top=Side(style='thin'),bottom=Side(style='thin'))
    hf = PatternFill("solid", fgColor="1F4E79")
    hfont = Font(name="Arial", bold=True, color="FFFFFF", size=11)
    dfont = Font(name="Arial", size=10)
    ctr = Alignment(horizontal="center", vertical="center", wrap_text=True)
    lft = Alignment(horizontal="left", vertical="center", wrap_text=True)
    
    ws.merge_cells("A1:H1")
    ws["A1"], ws["A1"].font, ws["A1"].alignment = "差旅发票汇总表", Font(name="Arial", bold=True, size=14), ctr
    headers = ["序号","类型","发票号码","行程日期","出发地/出发站","到达地/到达站","金额 (元)","Word 页码"]
    for col, h in enumerate(headers, 1):
        c = ws.cell(row=2, column=col, value=h)
        c.font, c.fill, c.alignment, c.border = hfont, hf, ctr, thin
    
    row = 3
    type_order = ['滴滴', '阳光出行', '通行费', '携程机票', '12306 火车票', '行程单']
    type_counters = {t: 0 for t in type_order}
    for inv_type in type_order:
        recs = [r for r in all_records if r.get('type') == inv_type]
        if not recs: continue
        for r in recs:
            type_counters[inv_type] += 1
            trip_count = r.get('trip_count', 0)
            trip_date = r.get('trip_date', '')
            from_loc = r.get('from', '详见发票')
            to_loc = r.get('to', '详见发票')
            if trip_count > 1:
                from_loc = f"对应{trip_count}个行程"
                to_loc = ""
            for col, v in enumerate([str(type_counters[inv_type]), r.get('type',''), r.get('inv_no',''), trip_date, from_loc, to_loc, r.get('amount',0), r.get('word_page','')], 1):
                c = ws.cell(row=row, column=col, value=v)
                c.font, c.border, c.alignment = dfont, thin, ctr if col in (1,2,3,4,7,8) else lft
                if col == 7: c.number_format = '#,##0.00'
            row += 1
    
    ws.cell(row=row, column=1, value="合计").font = Font(name="Arial", bold=True, color="FFFFFF", size=11)
    ws.cell(row=row, column=1).fill, ws.cell(row=row, column=1).alignment, ws.cell(row=row, column=1).border = hf, ctr, thin
    ws.merge_cells(f"A{row}:F{row}")
    total = round(sum(r.get('amount', 0) for r in all_records), 2)
    ws.cell(row=row, column=7, value=total)
    ws.cell(row=row, column=7).number_format, ws.cell(row=row, column=7).font = '#,##0.00', Font(name="Arial", bold=True, size=11)
    ws.cell(row=row, column=7).alignment, ws.cell(row=row, column=7).fill = ctr, PatternFill("solid", fgColor="D6E4F0")
    ws.cell(row=row, column=7).border = thin
    for i, w in enumerate([6,10,30,22,26,26,13,16], 1): ws.column_dimensions[chr(64+i)].width = w
    ws.row_dimensions[1].height, ws.row_dimensions[2].height = 24, 20
    for r in range(3, row+1): ws.row_dimensions[r].height = 36
    out = f'{OUTPUT_DIR}/差旅发票汇总.xlsx'
    wb.save(out)
    print(f"  已保存：{out}")
    return out

def main():
    print("=" * 50 + "\n  差旅发票整理工具\n" + "=" * 50)
    cfg = load_config()
    print(f"\n📅 请输入日期范围（格式：YYYY-MM-DD）")
    start = input(f"  开始日期（如 2026-04-01）: ").strip()
    end = input(f"  结束日期（如 2026-04-30）: ").strip()
    if not start or not end: print("日期不能为空，退出。"); sys.exit(1)
    invoice_emails = scan_invoices(cfg, start, end)
    if not invoice_emails: print("❌ 未找到任何发票邮件。"); sys.exit(0)
    all_records = download_and_process(cfg, invoice_emails)
    if not all_records: print("❌ 未下载到任何发票。"); sys.exit(1)
    build_word(all_records)
    build_excel(all_records)
    print(f"\n✅ 完成！共整理 {len(all_records)} 条记录")
    print(f"  📂 输出目录：{OUTPUT_DIR}")
    print(f"  📁 发票 PDF: {OUTPUT_DIR}/发票 PDF/ ({len([f for f in os.listdir(f'{OUTPUT_DIR}/发票 PDF') if f.endswith('.pdf')])}个)")
    print(f"  📁 行程单 PDF: {TRIP_SHEET_DIR} ({len([f for f in os.listdir(TRIP_SHEET_DIR) if f.endswith('.pdf')])}个)")

if __name__ == '__main__':
    main()
